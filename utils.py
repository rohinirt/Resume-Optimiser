import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import html
import re
from io import BytesIO
from pypdf import PdfReader

def extract_text_from_file(uploaded_file):
    if not uploaded_file:
        return ""
    text = ""
    uploaded_file.seek(0)
    if uploaded_file.name.endswith(".pdf"):
        pdf = PdfReader(uploaded_file)
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"
    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    uploaded_file.seek(0)
    return text

def add_bottom_border(paragraph, color_hex="0F172A", size="12"):
    """Adds a horizontal line directly under section headings."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def clean_markdown_bold_spans(text):
    """Safely parses markdown **bold** pairs to prevent trailing bold leaks."""
    if not text:
        return []
    
    pattern = re.compile(r'\*\*(.*?)\*\*')
    spans = []
    last_idx = 0
    
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            spans.append((text[last_idx:start], False))
        spans.append((match.group(1), True))
        last_idx = end
        
    if last_idx < len(text):
        remaining = text[last_idx:].replace("**", "")
        if remaining:
            spans.append((remaining, False))
            
    return spans

def render_spans_to_html(text):
    """Converts markdown **bold** spans into clean HTML <strong> tags."""
    spans = clean_markdown_bold_spans(text)
    out_html = ""
    for part, is_bold in spans:
        escaped_part = html.escape(part)
        if is_bold:
            out_html += f'<strong>{escaped_part}</strong>'
        else:
            out_html += escaped_part
    return out_html

def generate_standard_resume_sheet_html(title_header, content_text_or_bytes, is_docx_file=False):
    """
    Renders Original Master Resume with the EXACT same structure, bullet dots,
    bold category prefixes, and spacing as the Tailored Resume Preview.
    """
    if is_docx_file and isinstance(content_text_or_bytes, bytes) and len(content_text_or_bytes) > 0:
        try:
            doc = docx.Document(BytesIO(content_text_or_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception:
            lines = []
    else:
        lines = [line.strip() for line in str(content_text_or_bytes).split('\n') if line.strip()]

    if not lines:
        return "<p>No content found.</p>"

    cand_name = html.escape(lines[0])
    cand_details = html.escape(lines[1]) if len(lines) > 1 else ""
    body_lines = lines[2:] if len(lines) > 2 else lines[1:]

    paragraphs_html = ""
    current_section = ""
    in_bullet_list = False

    for txt in body_lines:
        escaped_txt = html.escape(txt)

        # 1. Detect Section Headers
        if txt.isupper() and len(txt) < 40:
            if in_bullet_list:
                paragraphs_html += "</ul>"
                in_bullet_list = False
            current_section = txt.upper()
            paragraphs_html += f'<div class="section-title">{escaped_txt}</div>'
            continue

        # 2. Format Bullet Points (if line starts with bullet symbols OR sits under Experience/Projects)
        is_bullet = txt.startswith("•") or txt.startswith("-") or txt.startswith("*") or (
            ("EXPERIENCE" in current_section or "PROJECTS" in current_section) and len(txt) > 30 and not any(k in txt for k in ["Jan 2", "Oct 2", "2026", "2025", "2024"])
        )

        if is_bullet:
            if not in_bullet_list:
                paragraphs_html += '<ul style="margin-top: 2px; margin-bottom: 8px; padding-left: 18px; font-size: 0.86rem; line-height: 1.5; color: #000000;">'
                in_bullet_list = True
            
            clean_bullet = txt.lstrip("•-* ").strip()
            formatted_bullet = render_spans_to_html(clean_bullet)
            paragraphs_html += f'<li style="margin-bottom: 4px; color: #000000;">{formatted_bullet}</li>'

        # 3. Format Non-Bullet Paragraphs (Technical Skills, Job Titles, Education)
        else:
            if in_bullet_list:
                paragraphs_html += "</ul>"
                in_bullet_list = False

            # Bold Skill Categories (e.g., "Visualization & BI:")
            if ":" in txt and ("SKILLS" in current_section or len(txt) < 80):
                parts = txt.split(":", 1)
                formatted_line = f'<strong>{html.escape(parts[0])}:</strong>{render_spans_to_html(parts[1])}'
                paragraphs_html += f'<div style="margin-bottom: 4px; font-size: 0.86rem; color: #000000;">{formatted_line}</div>'
            
            # Bold Role Titles (e.g., "Data Analytics Specialist | Uber")
            elif "EXPERIENCE" in current_section or "PROJECTS" in current_section:
                formatted_line = render_spans_to_html(txt)
                paragraphs_html += f'<p style="font-weight: 700; color: #000000; margin-bottom: 2px; font-size: 0.92rem; margin-top: 10px;">{formatted_line}</p>'
            
            # Education & Certifications
            elif "EDUCATION" in current_section or "CERTIFICATIONS" in current_section:
                if "," in txt:
                    parts = txt.split(",", 1)
                    paragraphs_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{html.escape(parts[0])}</strong>,{html.escape(parts[1])}</div>'
                else:
                    paragraphs_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{escaped_txt}</strong></div>'
            
            # Standard Body Paragraph (Summary)
            else:
                formatted_line = render_spans_to_html(txt)
                paragraphs_html += f'<p style="font-size: 0.86rem; line-height: 1.5; color: #000000; margin-bottom: 12px;">{formatted_line}</p>'

    if in_bullet_list:
        paragraphs_html += "</ul>"

    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, 'Helvetica Neue', Helvetica, sans-serif;
                background-color: #ffffff;
                color: #000000;
                margin: 0;
                padding: 25px;
            }}
            .header-name {{
                font-size: 1.4rem;
                font-weight: 800;
                color: #000000;
                text-align: center;
                letter-spacing: 0.5px;
            }}
            .header-contact {{
                font-size: 0.8rem;
                color: #000000;
                text-align: center;
                margin-top: 2px;
                margin-bottom: 14px;
            }}
            .section-title {{
                color: #1e3a8a;
                font-size: 10pt;
                margin-top: 12px;
                margin-bottom: 6px;
                border-bottom: 1.5px solid #0f172a;
                padding-bottom: 2px;
                font-weight: 700;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }}
        </style>
    </head>
    <body>
        <div class="header-name">{cand_name}</div>
        <div class="header-contact">{cand_details}</div>
        {paragraphs_html}
    </body>
    </html>
    """

def generate_paper_sheet_tailored_html(results):
    """Renders Tailored Resume HTML preview with properly formatted dates and selective education/cert bolding."""
    sec2 = results.get("section_2_tailored_content", {})
    keywords = results.get("post_optimization", {}).get("matching_keywords", [])
    
    contact = sec2.get("contact_info", {})
    cand_name = html.escape(str(contact.get("name", "ROHINI TEMBHURNIKAR")))
    cand_details = html.escape(str(contact.get("details", "(+91) 8010132326 | rohinitembhurnikar3@gmail.com | Hyderabad | LinkedIn | GitHub")))

    summary = html.escape(str(sec2.get("professional_summary", "")))
    skills_grouped = sec2.get("core_competencies_grouped", {})
    exp_list = sec2.get("professional_experience", [])
    proj_list = sec2.get("projects", [])
    edu_list = sec2.get("education", [])
    cert_list = sec2.get("certifications", [])

    for kw in keywords:
        if len(kw) > 2 and kw in summary:
            escaped_kw = html.escape(kw)
            summary = summary.replace(
                escaped_kw, 
                f'<mark style="background-color: #fef08a; text-decoration: underline; padding: 1px 4px; border-radius: 3px; font-weight: 600; color: #000000;">{escaped_kw}</mark>'
            )

    skills_html = ""
    for cat, val in skills_grouped.items():
        skills_html += f'<div style="margin-bottom: 4px; font-size: 0.86rem; color: #000000;"><strong>{html.escape(str(cat))}:</strong> <span style="color: #000000;">{html.escape(str(val))}</span></div>'

    exp_html = ""
    for role in exp_list:
        role_title = html.escape(str(role.get("role_title", "")))
        exp_html += f'<p style="font-weight: 700; color: #000000; margin-bottom: 2px; font-size: 0.92rem; margin-top: 10px;">{role_title}</p><ul style="margin-top: 2px; margin-bottom: 8px; padding-left: 18px; font-size: 0.86rem; line-height: 1.5; color: #000000;">'
        for b in role.get("bullets", []):
            formatted_b = render_spans_to_html(b)
            exp_html += f'<li style="margin-bottom: 4px; color: #000000;">{formatted_b}</li>'
        exp_html += '</ul>'

    proj_html = ""
    for proj in proj_list:
        proj_title = html.escape(str(proj.get("project_title", "")))
        proj_html += f'<p style="font-weight: 700; color: #000000; margin-bottom: 2px; font-size: 0.92rem; margin-top: 10px;">{proj_title}</p><ul style="margin-top: 2px; margin-bottom: 8px; padding-left: 18px; font-size: 0.86rem; line-height: 1.5; color: #000000;">'
        for b in proj.get("bullets", []):
            formatted_b = render_spans_to_html(b)
            proj_html += f'<li style="margin-bottom: 4px; color: #000000;">{formatted_b}</li>'
        proj_html += '</ul>'

    # Education: Bold ONLY degree name (part before comma or pipe)
    edu_html = ""
    for e in edu_list:
        txt = str(e)
        if "," in txt:
            parts = txt.split(",", 1)
            edu_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{html.escape(parts[0].strip())}</strong>, {html.escape(parts[1].strip())}</div>'
        elif "|" in txt:
            parts = txt.split("|", 1)
            edu_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{html.escape(parts[0].strip())}</strong> | {html.escape(parts[1].strip())}</div>'
        else:
            edu_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;">{html.escape(txt)}</div>'

    # Certifications: Bold ONLY certification name
    cert_html = ""
    for c in cert_list:
        txt = str(c)
        if "," in txt:
            parts = txt.split(",", 1)
            cert_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{html.escape(parts[0].strip())}</strong>, {html.escape(parts[1].strip())}</div>'
        elif "|" in txt:
            parts = txt.split("|", 1)
            cert_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;"><strong style="font-size: 9.5pt;">{html.escape(parts[0].strip())}</strong> | {html.escape(parts[1].strip())}</div>'
        else:
            cert_html += f'<div style="font-size: 0.86rem; margin-bottom: 3px; color: #000000;">{html.escape(txt)}</div>'

    full_document_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, 'Helvetica Neue', Helvetica, sans-serif;
                background-color: #ffffff;
                color: #000000;
                margin: 0;
                padding: 25px;
            }}
            .header-name {{
                font-size: 1.4rem;
                font-weight: 800;
                color: #000000;
                text-align: center;
                letter-spacing: 0.5px;
            }}
            .header-contact {{
                font-size: 0.8rem;
                color: #000000;
                text-align: center;
                margin-top: 2px;
                margin-bottom: 14px;
            }}
            .section-title {{
                color: #1e3a8a;
                font-size: 10pt;
                margin-top: 12px;
                margin-bottom: 6px;
                border-bottom: 1.5px solid #0f172a;
                padding-bottom: 2px;
                font-weight: 700;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }}
        </style>
    </head>
    <body>
        <div class="header-name">{cand_name}</div>
        <div class="header-contact">{cand_details}</div>

        <div class="section-title">Professional Summary</div>
        <p style="font-size: 0.86rem; line-height: 1.5; color: #000000; margin-bottom: 12px;">{summary}</p>

        <div class="section-title">Technical Skills</div>
        <div>{skills_html}</div>

        <div class="section-title">Work Experience</div>
        <div>{exp_html}</div>

        <div class="section-title">Projects</div>
        <div>{proj_html}</div>

        <div class="section-title">Education</div>
        <div>{edu_html}</div>

        <div class="section-title">Certifications</div>
        <div>{cert_html}</div>
    </body>
    </html>
    """
    return full_document_html

def generate_new_formatted_docx(results):
    """Generates 1-Page A4 Word document with exact selective bolding for Education/Certifications."""
    output = BytesIO()
    doc = docx.Document()

    sec2 = results.get("section_2_tailored_content", {})

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Header
    contact = sec2.get("contact_info", {})
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    r_name = p_name.add_run(contact.get("name", "ROHINI TEMBHURNIKAR"))
    r_name.font.name = "Arial"
    r_name.font.size = Pt(15)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(0, 0, 0)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(8)
    r_contact = p_contact.add_run(contact.get("details", ""))
    r_contact.font.name = "Arial"
    r_contact.font.size = Pt(8.5)
    r_contact.font.color.rgb = RGBColor(0, 0, 0)

    def add_section_header(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title_text.upper())
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        add_bottom_border(p, color_hex="0F172A", size="8")
        return p

    # Summary
    add_section_header("PROFESSIONAL SUMMARY")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(6)
    
    sum_text = sec2.get("professional_summary", "")
    sum_spans = clean_markdown_bold_spans(sum_text)
    for part, is_bold in sum_spans:
        r_sum = p_sum.add_run(part)
        r_sum.font.name = "Arial"
        r_sum.font.size = Pt(9)
        r_sum.font.bold = is_bold
        r_sum.font.color.rgb = RGBColor(0, 0, 0)

    # Technical Skills
    add_section_header("TECHNICAL SKILLS")
    grouped_skills = sec2.get("core_competencies_grouped", {})
    for cat, val in grouped_skills.items():
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(0)
        p_sk.paragraph_format.space_after = Pt(2)
        
        r_cat = p_sk.add_run(f"{cat}: ")
        r_cat.font.name = "Arial"
        r_cat.font.size = Pt(9)
        r_cat.font.bold = True
        r_cat.font.color.rgb = RGBColor(0, 0, 0)

        r_val = p_sk.add_run(str(val))
        r_val.font.name = "Arial"
        r_val.font.size = Pt(9)
        r_val.font.bold = False
        r_val.font.color.rgb = RGBColor(0, 0, 0)

    # Work Experience
    add_section_header("WORK EXPERIENCE")
    for role in sec2.get("professional_experience", []):
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(4)
        p_role.paragraph_format.space_after = Pt(2)
        
        r_role = p_role.add_run(role.get("role_title", "Role"))
        r_role.font.name = "Arial"
        r_role.font.size = Pt(9.5)
        r_role.font.bold = True
        r_role.font.color.rgb = RGBColor(0, 0, 0)

        for b in role.get("bullets", []):
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            
            b_spans = clean_markdown_bold_spans(b.strip())
            for part, is_bold in b_spans:
                r_b = p_b.add_run(part)
                r_b.font.name = "Arial"
                r_b.font.size = Pt(8.8)
                r_b.font.bold = is_bold
                r_b.font.color.rgb = RGBColor(0, 0, 0)

    # Projects
    add_section_header("PROJECTS")
    for proj in sec2.get("projects", []):
        p_proj = doc.add_paragraph()
        p_proj.paragraph_format.space_before = Pt(4)
        p_proj.paragraph_format.space_after = Pt(2)
        
        r_proj = p_proj.add_run(proj.get("project_title", "Project"))
        r_proj.font.name = "Arial"
        r_proj.font.size = Pt(9.5)
        r_proj.font.bold = True
        r_proj.font.color.rgb = RGBColor(0, 0, 0)

        for b in proj.get("bullets", []):
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            
            p_spans = clean_markdown_bold_spans(b.strip())
            for part, is_bold in p_spans:
                r_b = p_b.add_run(part)
                r_b.font.name = "Arial"
                r_b.font.size = Pt(8.8)
                r_b.font.bold = is_bold
                r_b.font.color.rgb = RGBColor(0, 0, 0)

    # Education (Fix: Only bold Degree Name at 9.5pt, rest regular text at 8.8pt)
    add_section_header("EDUCATION")
    for edu in sec2.get("education", []):
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(0)
        p_edu.paragraph_format.space_after = Pt(2)
        
        txt = str(edu)
        if "," in txt:
            parts = txt.split(",", 1)
            r_deg = p_edu.add_run(parts[0].strip())
            r_deg.font.name = "Arial"
            r_deg.font.size = Pt(9.5)
            r_deg.font.bold = True
            r_deg.font.color.rgb = RGBColor(0, 0, 0)
            
            r_rest = p_edu.add_run(f", {parts[1].strip()}")
            r_rest.font.name = "Arial"
            r_rest.font.size = Pt(8.8)
            r_rest.font.bold = False
            r_rest.font.color.rgb = RGBColor(0, 0, 0)
        elif "|" in txt:
            parts = txt.split("|", 1)
            r_deg = p_edu.add_run(parts[0].strip())
            r_deg.font.name = "Arial"
            r_deg.font.size = Pt(9.5)
            r_deg.font.bold = True
            r_deg.font.color.rgb = RGBColor(0, 0, 0)
            
            r_rest = p_edu.add_run(f" | {parts[1].strip()}")
            r_rest.font.name = "Arial"
            r_rest.font.size = Pt(8.8)
            r_rest.font.bold = False
            r_rest.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r_deg = p_edu.add_run(txt)
            r_deg.font.name = "Arial"
            r_deg.font.size = Pt(8.8)
            r_deg.font.bold = False
            r_deg.font.color.rgb = RGBColor(0, 0, 0)

    # Certifications (Fix: Only bold Certificate Name at 9.5pt, rest regular text at 8.8pt)
    add_section_header("CERTIFICATIONS")
    for cert in sec2.get("certifications", []):
        p_cert = doc.add_paragraph()
        p_cert.paragraph_format.space_before = Pt(0)
        p_cert.paragraph_format.space_after = Pt(2)
        
        txt = str(cert)
        if "," in txt:
            parts = txt.split(",", 1)
            r_cert = p_cert.add_run(parts[0].strip())
            r_cert.font.name = "Arial"
            r_cert.font.size = Pt(9.5)
            r_cert.font.bold = True
            r_cert.font.color.rgb = RGBColor(0, 0, 0)
            
            r_rest = p_cert.add_run(f", {parts[1].strip()}")
            r_rest.font.name = "Arial"
            r_rest.font.size = Pt(8.8)
            r_rest.font.bold = False
            r_rest.font.color.rgb = RGBColor(0, 0, 0)
        elif "|" in txt:
            parts = txt.split("|", 1)
            r_cert = p_cert.add_run(parts[0].strip())
            r_cert.font.name = "Arial"
            r_cert.font.size = Pt(9.5)
            r_cert.font.bold = True
            r_cert.font.color.rgb = RGBColor(0, 0, 0)
            
            r_rest = p_cert.add_run(f" | {parts[1].strip()}")
            r_rest.font.name = "Arial"
            r_rest.font.size = Pt(8.8)
            r_rest.font.bold = False
            r_rest.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r_cert = p_cert.add_run(txt)
            r_cert.font.name = "Arial"
            r_cert.font.size = Pt(8.8)
            r_cert.font.bold = False
            r_cert.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(output)
    output.seek(0)
    return output

